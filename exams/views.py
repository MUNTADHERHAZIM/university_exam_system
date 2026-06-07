import json
import random
from datetime import datetime
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib import messages
from django.http import JsonResponse, HttpResponseForbidden
from django.utils import timezone
from django.views.decorators.http import require_POST
from django.db.models import Avg, Count, Q
from .models import (
    Course, Exam, Question, QuestionOption, QuestionBlank, MatchingPair,
    OrderingItem, ExamAttempt, AttemptAnswer, ViolationLog, Notification,
    ChatMessage, ContactMessage
)
from .grading import auto_grade_attempt
from accounts.models import Profile
import csv
from django.http import HttpResponse
from django.db.models import F, Max, Min
from .models import ProctorSnapshot
import openpyxl
from docx import Document


def require_role(roles):
    if isinstance(roles, str):
        roles = [roles]
    def decorator(view_func):
        @login_required
        def wrapper(request, *args, **kwargs):
            try:
                if request.user.profile.role not in roles:
                    return render(request, 'errors/403.html', status=403)
            except Profile.DoesNotExist:
                return render(request, 'errors/403.html', status=403)
            return view_func(request, *args, **kwargs)
        return wrapper
    return decorator


@login_required
def dashboard(request):
    profile, created = Profile.objects.get_or_create(user=request.user)
    if created and (request.user.is_superuser or request.user.is_staff):
        profile.role = 'admin'
        profile.save()

    # Determine the view based on the ACTIVE ROLE from profile
    active_role = profile.role

    if active_role in ['admin', 'assistant']:
        # ADMIN / ASSISTANT (Proctor/Grader) DASHBOARD
        if request.user.is_superuser:
            # Superuser sees all exams and attempts
            exams = Exam.objects.all()
            total_attempts = ExamAttempt.objects.filter(is_submitted=True)
        else:
            # Admins and Assistants see exams they created or all if they are assistant? 
            # Let's show all for Assistant, and only owned for Admin on the MAIN dashboard (stats)
            if active_role == 'assistant':
                exams = Exam.objects.all()
                total_attempts = ExamAttempt.objects.filter(is_submitted=True)
            else:
                exams = Exam.objects.filter(created_by=request.user)
                total_attempts = ExamAttempt.objects.filter(exam__created_by=request.user, is_submitted=True)
            
        pending_grading = total_attempts.filter(is_fully_graded=False).count()
        students_count = User.objects.filter(profile__role='student').count()
        recent_attempts = total_attempts.order_by('-submitted_at')[:5]
        
        context = {
            'profile': profile,
            'exams': exams,
            'total_exams': exams.count(),
            'total_attempts': total_attempts.count(),
            'pending_grading': pending_grading,
            'students_count': students_count,
            'recent_attempts': recent_attempts,
            'active_exams': exams.filter(status='active').count(),
            'contact_messages': ContactMessage.objects.filter(is_resolved=False).order_by('-created_at')[:5],
        }
        return render(request, 'exams/dashboard_admin.html', context)
    else:
        # STUDENT DASHBOARD
        # Filter exams based on targeting:
        # 1. Active status
        # 2. (No restrictions OR specifically assigned to student OR student's college matches OR student's department matches)
        user_college = profile.college
        user_dept = profile.academic_department
        
        available_exams = Exam.objects.filter(status='active').filter(
            Q(assigned_students__isnull=True, assigned_colleges__isnull=True, assigned_departments__isnull=True) |
            Q(assigned_students=request.user) |
            Q(assigned_colleges=user_college) |
            Q(assigned_departments=user_dept)
        ).distinct()
        
        my_attempts = ExamAttempt.objects.filter(student=request.user, is_submitted=True).select_related('exam')
        my_attempts = ExamAttempt.objects.filter(student=request.user, is_submitted=True).select_related('exam').order_by('-submitted_at')
        
        # Calculate attempts count and max allowed for each available exam
        for exam in available_exams:
            used = my_attempts.filter(exam=exam).count()
            override = exam.student_overrides.filter(student=request.user).first()
            max_allowed = exam.max_attempts + (override.extra_attempts if override else 0)
            
            exam.used_attempts = used
            exam.max_allowed = max_allowed
            exam.remaining_attempts = max_allowed - used
            exam.can_retake = used < max_allowed

        context = {
            'available_exams': available_exams,
            'my_attempts': my_attempts,
            'passed_count': sum(1 for a in my_attempts if a.get_percentage() >= (a.exam.pass_mark / a.exam.total_marks * 100 if a.exam.total_marks else 0)),
        }
        return render(request, 'exams/dashboard_student.html', context)


# ─── EXAM MANAGEMENT (Admin) ─────────────────────────────────────────────────

@require_role('admin')
def exam_list(request):
    exams = Exam.objects.filter(created_by=request.user).annotate(
        attempt_count=Count('attempts', filter=Q(attempts__is_submitted=True))
    )
    # Search
    search_q = request.GET.get('q', '').strip()
    if search_q:
        exams = exams.filter(Q(title__icontains=search_q) | Q(subject__icontains=search_q))
    # Subject filter
    subject_filter = request.GET.get('subject', '')
    if subject_filter:
        exams = exams.filter(subject=subject_filter)
    # Status filter
    status_filter = request.GET.get('status', '')
    if status_filter:
        exams = exams.filter(status=status_filter)

    # Academic filters
    college_id = request.GET.get('college')
    dept_id = request.GET.get('department')
    if college_id:
        exams = exams.filter(course__department__college_id=college_id)
    if dept_id:
        exams = exams.filter(course__department_id=dept_id)

    # Fetch choices for filters
    from accounts.models import College, Department
    colleges = College.objects.all()
    departments = Department.objects.all()
    if college_id:
        departments = departments.filter(college_id=college_id)

    subjects = Exam.objects.filter(created_by=request.user).values_list('subject', flat=True).distinct()
    context = {
        'exams': exams,
        'subjects': subjects,
        'search_q': search_q,
        'subject_filter': subject_filter,
        'status_filter': status_filter,
        'colleges': colleges,
        'departments': departments,
        'selected_college': int(college_id) if college_id else None,
        'selected_dept': int(dept_id) if dept_id else None,
    }
    return render(request, 'exams/exam_list.html', context)


@require_role(['admin', 'assistant'])
def exam_create(request):
    from accounts.models import College, Department
    if request.method == 'POST':
        exam = Exam.objects.create(
            title=request.POST['title'],
            subject=request.POST['subject'],
            course_id=request.POST.get('course'),
            description=request.POST.get('description', ''),
            duration=int(request.POST.get('duration', 90)),
            total_marks=int(request.POST.get('total_marks', 100)),
            pass_mark=int(request.POST.get('pass_mark', 50)),
            max_attempts=int(request.POST.get('max_attempts', 1)),
            shuffle_questions=request.POST.get('shuffle_questions') == 'on',
            allow_review=request.POST.get('allow_review') == 'on',
            status='draft',
            created_by=request.user,
        )
        messages.success(request, f'تم إنشاء الاختبار "{exam.title}" بنجاح')
        return redirect('exam_edit', pk=exam.pk)
    
    colleges = College.objects.all()
    departments = Department.objects.all()
    courses = Course.objects.all()
    
    return render(request, 'exams/exam_create.html', {
        'colleges': colleges,
        'departments': departments,
        'courses': courses,
    })


@require_role(['admin', 'assistant'])
def exam_edit(request, pk):
    from accounts.models import College, Department
    exam = get_object_or_404(Exam, pk=pk)
    
    # Permission: Admins edit everything, others only their own
    if not request.user.profile.is_admin() and exam.created_by != request.user:
        return render(request, 'errors/403.html', status=403)
    if request.method == 'POST':
        exam.title = request.POST.get('title', exam.title)
        exam.subject = request.POST.get('subject', exam.subject)
        exam.course_id = request.POST.get('course') or None
        exam.description = request.POST.get('description', exam.description)
        exam.duration = int(request.POST.get('duration', exam.duration))
        exam.total_marks = int(request.POST.get('total_marks', exam.total_marks))
        exam.pass_mark = int(request.POST.get('pass_mark', exam.pass_mark))
        exam.max_attempts = int(request.POST.get('max_attempts', exam.max_attempts))
        exam.status = request.POST.get('status', exam.status)
        exam.shuffle_questions = request.POST.get('shuffle_questions') == 'on'
        exam.allow_review = request.POST.get('allow_review') == 'on'
        exam.is_mock = request.POST.get('is_mock') == 'on'
        exam.require_seb = request.POST.get('require_seb') == 'on'
        exam.require_camera = request.POST.get('require_camera') == 'on'
        
        # Advanced Features
        random_count = request.POST.get('random_question_count')
        exam.random_question_count = int(random_count) if random_count else None
        
        # Handle scheduling dates
        start_date_str = request.POST.get('start_date', '')
        end_date_str = request.POST.get('end_date', '')
        exam.start_date = datetime.fromisoformat(start_date_str) if start_date_str else None
        exam.end_date = datetime.fromisoformat(end_date_str) if end_date_str else None
        old_status = Exam.objects.get(pk=pk).status
        exam.save()
        
        # Targeting
        assigned_ids = request.POST.getlist('assigned_students')
        exam.assigned_students.set(assigned_ids) if assigned_ids else exam.assigned_students.clear()
        
        college_ids = request.POST.getlist('assigned_colleges')
        exam.assigned_colleges.set(college_ids) if college_ids else exam.assigned_colleges.clear()
        
        dept_ids = request.POST.getlist('assigned_departments')
        exam.assigned_departments.set(dept_ids) if dept_ids else exam.assigned_departments.clear()

        # Notify students when exam becomes active
        if old_status != 'active' and exam.status == 'active':
            students = User.objects.filter(profile__role='student')
            for student in students:
                create_notification(
                    user=student,
                    notification_type='new_exam',
                    title=f'اختبار جديد: {exam.title}',
                    message=f'تم نشر اختبار جديد في مادة {exam.subject}',
                    link=f'/exams/{exam.pk}/start/',
                )
        # Final Sync Check: If total_marks is 0 but there are questions, or if 'sync_marks' was requested
        q_sum = exam.get_questions_total_marks()
        if request.POST.get('sync_marks') == 'true' or (exam.total_marks == 0 and q_sum > 0):
            exam.recalculate_total_marks()
            messages.info(request, f'تمت موازنة الدرجة الكلية لتصبح {exam.total_marks} بناءً على مجموع درجات الأسئلة.')

        messages.success(request, 'تم حفظ التعديلات وحفظ الإعدادات المتقدمة بنجاح')
        return redirect('exam_edit', pk=exam.pk)

    questions = exam.questions.prefetch_related('options', 'blanks', 'pairs', 'order_items')
    all_students = User.objects.filter(profile__role='student').order_by('first_name', 'username')
    assigned_ids = list(exam.assigned_students.values_list('id', flat=True))
    overrides = exam.student_overrides.select_related('student')
    
    colleges = College.objects.all()
    departments = Department.objects.all()
    courses = Course.objects.all()
    questions_total = exam.get_questions_total_marks()
    
    context = {
        'exam': exam,
        'questions': questions,
        'all_students': all_students,
        'assigned_ids': assigned_ids,
        'overrides': overrides,
        'colleges': colleges,
        'departments': departments,
        'courses': courses,
        'questions_total': questions_total,
    }
    return render(request, 'exams/exam_edit.html', context)

def about_system(request):
    faqs = [
        {
            'q': 'كيف يمكنني البدء في استخدام النظام؟',
            'a': 'يمكنك البدء بإنشاء حساب طالب أو التواصل مع الإدارة لإعطائك صلاحيات أستاذ لإنشاء الاختبارات.'
        },
        {
            'q': 'هل يدعم النظام الهواتف المحمولة؟',
            'a': 'نعم، النظام مصمم بشكل متجاوب بالكامل ليعمل على كافة الشاشات والأجهزة اللوحية.'
        },
        {
            'q': 'كيف يتم ضمان نزاهة الاختبارات؟',
            'a': 'يعتمد النظام على تقنيات مراقبة تفاعلية مثل كشف محاولات الغش، تسجيل المخالفات تلقائياً، ومنع الخروج من المتصفح.'
        },
    ]
    return render(request, 'exams/about_system.html', {'faqs': faqs})


@require_role('admin')
def exam_live_monitor(request, pk):
    exam = get_object_or_404(Exam, pk=pk, created_by=request.user)
    return render(request, 'exams/exam_live_monitor.html', {'exam': exam})


@require_role('admin')
def api_live_monitor(request, pk):
    exam = get_object_or_404(Exam, pk=pk, created_by=request.user)
    
    recent_limit = timezone.now() - timezone.timedelta(hours=exam.duration)
    active_attempts = ExamAttempt.objects.filter(
        exam=exam,
        started_at__gte=recent_limit
    ).select_related('student', 'exam').order_by('-started_at')

    data = []
    now = timezone.now()
    for attempt in active_attempts:
        override = exam.student_overrides.filter(student=attempt.student).first()
        extra_time = override.extra_time_minutes if override else 0
        total_time_seconds = (exam.duration + extra_time) * 60
        elapsed = (now - attempt.started_at).total_seconds()
        
        if attempt.is_submitted:
            status = 'submitted'
            time_left_str = 'منتهي'
            pct = attempt.get_percentage()
        else:
            time_remaining = max(0, int(total_time_seconds - elapsed))
            if time_remaining == 0:
                status = 'time_up'
            else:
                status = 'active'
            
            m, s = divmod(time_remaining, 60)
            time_left_str = f"{m}د {s}ث"
            pct = None

        latest_snap = attempt.snapshots.first()
        
        data.append({
            'student_name': attempt.student.get_full_name() or attempt.student.username,
            'status': status,
            'violations': attempt.violations_count,
            'time_left': time_left_str,
            'started_at': attempt.started_at.strftime('%H:%M'),
            'score_pct': pct,
            'ip': attempt.ip_address,
            'attempt_id': attempt.id,
            'latest_snapshot': latest_snap.image.url if latest_snap else None,
        })

    return JsonResponse({'attempts': data})


@require_role('admin')
@require_POST
def api_save_override(request, pk):
    exam = get_object_or_404(Exam, pk=pk, created_by=request.user)
    try:
        data = json.loads(request.body)
        student_id = data.get('student_id')
        extra_attempts = int(data.get('extra_attempts', 0))
        extra_time = int(data.get('extra_time_minutes', 0))
        
        student = get_object_or_404(User, pk=student_id, profile__role='student')
        
        from .models import StudentExamOverride
        override, created = StudentExamOverride.objects.update_or_create(
            exam=exam,
            student=student,
            defaults={
                'extra_attempts': extra_attempts,
                'extra_time_minutes': extra_time
            }
        )
        return JsonResponse({'status': 'success', 'msg': 'تم حفظ الاستثناء بنجاح'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'msg': str(e)}, status=400)


@require_role(['admin', 'assistant'])
def exam_statistics(request, pk):
    exam = get_object_or_404(Exam, pk=pk)
    # Access control
    profile = getattr(request.user, 'profile', None)
    is_staff = (profile and profile.role in ['admin', 'assistant']) or request.user.is_superuser
    if not is_staff and exam.created_by != request.user:
        return HttpResponseForbidden()
    
    attempts = exam.attempts.filter(is_submitted=True)
    total_attempts = attempts.count()
    
    if total_attempts == 0:
        return render(request, 'exams/exam_statistics.html', {
            'exam': exam,
            'no_data': True
        })

    # Basic stats
    avg_score = attempts.aggregate(Avg('final_score'))['final_score__avg'] or 0
    max_score = attempts.aggregate(Max('final_score'))['final_score__max'] or 0
    min_score = attempts.aggregate(Min('final_score'))['final_score__min'] or 0
    
    pass_mark = exam.pass_mark
    passed_count = attempts.filter(final_score__gte=pass_mark).count()
    failed_count = total_attempts - passed_count
    
    # Score distribution for chart
    # Groups: 0-20%, 21-40%, 41-60%, 61-80%, 81-100%
    distribution = [0, 0, 0, 0, 0]
    total_time = 0
    for a in attempts:
        pct = a.get_percentage()
        total_time += a.time_spent_seconds
        if pct <= 20: distribution[0] += 1
        elif pct <= 40: distribution[1] += 1
        elif pct <= 60: distribution[2] += 1
        elif pct <= 80: distribution[3] += 1
        else: distribution[4] += 1

    avg_time_seconds = total_time / total_attempts
    m, s = divmod(int(avg_time_seconds), 60)
    avg_time_str = f"{m}د {s}ث"

    # Question Difficulty Analysis
    questions_data = []
    for q in exam.questions.all().order_by('order'):
        # Correctly join to AttemptAnswer using the related_name 'instances' as defined in models.py fix
        total_answers = AttemptAnswer.objects.filter(question=q, attempt__is_submitted=True).count()
        if total_answers > 0:
            # Consider 50% or more of the marks as "success" for that specific question
            correct_answers = AttemptAnswer.objects.filter(
                question=q, 
                attempt__is_submitted=True,
                earned_marks__gte=F('question__marks') * 0.5
            ).count()
            success_rate = round((correct_answers / total_answers) * 100, 1)
        else:
            success_rate = 0
            
        questions_data.append({
            'id': q.id,
            'text': q.text[:80],
            'type': q.get_question_type_display(),
            'success_rate': success_rate,
            'color': 'success' if success_rate > 70 else ('warning' if success_rate > 40 else 'danger')
        })

    context = {
        'exam': exam,
        'total_attempts': total_attempts,
        'avg_score': round(avg_score, 1),
        'max_score': max_score,
        'min_score': min_score,
        'avg_time_str': avg_time_str,
        'passed_count': passed_count,
        'failed_count': failed_count,
        'distribution': json.dumps(distribution),
        'questions_data': questions_data,
    }
    return render(request, 'exams/exam_statistics.html', context)


@require_role(['admin', 'assistant'])
def exam_delete(request, pk):
    exam = get_object_or_404(Exam, pk=pk)
    
    if request.method == 'POST':
        title = exam.title
        exam.delete()
        messages.success(request, f'تم حذف الاختبار "{title}"')
        return redirect('exam_list')
    return render(request, 'exams/exam_confirm_delete.html', {'exam': exam})


# ─── QUESTION MANAGEMENT ─────────────────────────────────────────────────────

def _save_question_components(question, request):
    qtype = question.question_type
    
    # Handle True/False
    if qtype == 'true_false':
        question.tf_answer = request.POST.get('tf_answer') == 'true'
        question.save()

    # Handle MCQ options
    elif qtype in ['mcq_single', 'mcq_multi']:
        question.options.all().delete()
        options = request.POST.getlist('option_text')
        corrects = request.POST.getlist('option_correct')
        for i, opt_text in enumerate(options):
            if opt_text.strip():
                QuestionOption.objects.create(
                    question=question,
                    text=opt_text.strip(),
                    is_correct=(str(i) in corrects),
                    order=i,
                )

    # Handle Fill blank
    elif qtype == 'fill_blank':
        question.blanks.all().delete()
        answers = request.POST.getlist('blank_answer')
        for ans in answers:
            if ans.strip():
                QuestionBlank.objects.create(
                    question=question,
                    accepted_answer=ans.strip(),
                    case_sensitive=False,
                )

    # Handle Matching
    elif qtype == 'matching':
        question.pairs.all().delete()
        lefts = request.POST.getlist('left_text')
        rights = request.POST.getlist('right_text')
        for i, (l, r) in enumerate(zip(lefts, rights)):
            if l.strip() and r.strip():
                MatchingPair.objects.create(
                    question=question,
                    left_text=l.strip(),
                    right_text=r.strip(),
                    order=i,
                )

    # Handle Ordering
    elif qtype == 'ordering':
        question.order_items.all().delete()
        items = request.POST.getlist('order_item')
        for i, item in enumerate(items):
            if item.strip():
                OrderingItem.objects.create(
                    question=question,
                    text=item.strip(),
                    correct_position=i,
                )

    # Handle Code
    elif qtype == 'code':
        question.code_language = request.POST.get('code_language', 'python')
        question.code_template = request.POST.get('code_template', '')
        question.save()


@require_role(['admin', 'assistant'])
def question_add(request, exam_pk):
    exam = get_object_or_404(Exam, pk=exam_pk)
    
    if not request.user.profile.is_admin() and exam.created_by != request.user:
        return HttpResponseForbidden("غير مصرح لك بإضافة أسئلة لهذا الاختبار")
    if request.method == 'POST':
        qtype = request.POST['question_type']
        marks = int(request.POST.get('marks', 5))
        order = exam.questions.count() + 1

        question = Question.objects.create(
            exam=exam,
            question_type=qtype,
            text=request.POST['text'],
            marks=marks,
            order=order,
            explanation=request.POST.get('explanation', ''),
            rubric=request.POST.get('rubric', ''),
            image=request.FILES.get('image'),
        )
        _save_question_components(question, request)
        messages.success(request, 'تمت إضافة السؤال بنجاح')
        return redirect('exam_edit', pk=exam_pk)

    return render(request, 'exams/question_form.html', {'exam': exam, 'question_types': Question.TYPE_CHOICES})


@require_role('admin')
def question_edit(request, pk):
    question = get_object_or_404(Question, pk=pk, exam__created_by=request.user)
    exam = question.exam
    if request.method == 'POST':
        question.question_type = request.POST['question_type']
        question.text = request.POST['text']
        question.marks = int(request.POST.get('marks', question.marks))
        question.explanation = request.POST.get('explanation', '')
        question.rubric = request.POST.get('rubric', '')
        
        if 'image' in request.FILES:
            question.image = request.FILES['image']
        elif 'clear_image' in request.POST:
            question.image = None
            
        question.save()
        
        _save_question_components(question, request)
        messages.success(request, 'تم تحديث السؤال بنجاح')
        return redirect('exam_edit', pk=exam.pk)

    return render(request, 'exams/question_form.html', {
        'exam': exam, 
        'question': question,
        'question_types': Question.TYPE_CHOICES
    })


@require_role('admin')
def question_delete(request, pk):
    question = get_object_or_404(Question, pk=pk)
    exam_pk = question.exam_id
    if request.method == 'POST':
        question.delete()
        messages.success(request, 'تم حذف السؤال')
    return redirect('exam_edit', pk=exam_pk)


# ─── STUDENT: TAKE EXAM ──────────────────────────────────────────────────────

@login_required
def exam_start(request, pk):
    exam = get_object_or_404(Exam, pk=pk)

    try:
        profile = request.user.profile
        if profile.is_admin():
            messages.warning(request, 'المشرفون لا يمكنهم أداء الاختبارات')
            return redirect('dashboard')
    except Profile.DoesNotExist:
        pass

    if not exam.is_available():
        messages.error(request, 'هذا الاختبار غير متاح حالياً')
        return redirect('dashboard')

    if exam.assigned_students.exists() and request.user not in exam.assigned_students.all():
        messages.error(request, 'غير مصرح لك بدخول هذا الاختبار. الاختبار مخصص لطلاب محددين.')
        return redirect('dashboard')

    # Check attempt limits and overrides
    override = exam.student_overrides.filter(student=request.user).first()
    max_allowed = exam.max_attempts + (override.extra_attempts if override else 0)
    attempts_count = ExamAttempt.objects.filter(exam=exam, student=request.user, is_submitted=True).count()
    
    if not exam.is_mock and attempts_count >= max_allowed:
        messages.info(request, f'لقد استنفدت جميع المحاولات المتاحة لهذا الاختبار ({max_allowed})')
        # Redirect to the latest result
        latest = ExamAttempt.objects.filter(exam=exam, student=request.user, is_submitted=True).first()
        if latest:
            return redirect('exam_result', pk=latest.pk)
        return redirect('dashboard')

    if request.method == 'POST':
        # Create a new attempt
        attempt = ExamAttempt.objects.create(
            exam=exam,
            student=request.user,
            ip_address=request.META.get('REMOTE_ADDR'),
            user_agent=request.META.get('HTTP_USER_AGENT', '')[:500],
        )

        # Generate random question bank subset
        qs = list(exam.questions.all())
        if exam.random_question_count and exam.random_question_count < len(qs):
            qs = random.sample(qs, exam.random_question_count)
        
        # Pre-create blank AttemptAnswers for the selected questions
        for q in qs:
            AttemptAnswer.objects.create(attempt=attempt, question=q)
            
        # Facial Recognition Photo
        snapshot_data = request.POST.get('snapshot_data')
        if snapshot_data and exam.require_camera:
            import base64
            from django.core.files.base import ContentFile
            try:
                format, imgstr = snapshot_data.split(';base64,')
                ext = format.split('/')[-1]
                data = ContentFile(base64.b64decode(imgstr), name=f'identity_{attempt.id}.{ext}')
                ProctorSnapshot.objects.create(attempt=attempt, image=data)
            except Exception:
                pass

        return redirect('exam_take', pk=attempt.pk)

    context = {
        'exam': exam,
        'attempts_used': attempts_count,
        'attempts_remaining': max_allowed - attempts_count,
    }
    return render(request, 'exams/exam_start.html', context)


@login_required
def exam_take(request, pk):
    attempt = get_object_or_404(ExamAttempt, pk=pk, student=request.user)

    if attempt.is_submitted:
        return redirect('exam_result', pk=attempt.pk)

    exam = attempt.exam
    
    if exam.require_seb:
        user_agent = request.META.get('HTTP_USER_AGENT', '')
        if 'SafeExamBrowser' not in user_agent:
            messages.error(request, 'هذا الاختبار يتطلب الدخول عبر متصفح Safe Exam Browser (SEB).')
            return redirect('dashboard')
    
    # Check assigned questions from Bank
    assigned_q_ids = attempt.answers.values_list('question_id', flat=True)
    if assigned_q_ids:
        questions = list(Question.objects.filter(id__in=assigned_q_ids).prefetch_related('options', 'blanks', 'pairs', 'order_items'))
    else:
        questions = list(exam.questions.prefetch_related('options', 'blanks', 'pairs', 'order_items'))

    if exam.shuffle_questions:
        random.shuffle(questions)

    # Shuffle options for MCQ
    for q in questions:
        if q.question_type in ['mcq_single', 'mcq_multi']:
            q._shuffled_options = list(q.options.all())
            random.shuffle(q._shuffled_options)
        if q.question_type == 'matching':
            q._shuffled_rights = list(q.pairs.all())
            random.shuffle(q._shuffled_rights)
        if q.question_type == 'ordering':
            q._shuffled_items = list(q.order_items.all())
            random.shuffle(q._shuffled_items)

    # Get existing answers
    existing_answers = {a.question_id: a for a in attempt.answers.all()}

    # Time remaining with overrides
    override = exam.student_overrides.filter(student=request.user).first()
    extra_time = override.extra_time_minutes if override else 0
    total_time_seconds = (exam.duration + extra_time) * 60

    elapsed = (timezone.now() - attempt.started_at).total_seconds()
    time_remaining = max(0, int(total_time_seconds - elapsed))

    if time_remaining == 0 and not attempt.is_submitted:
        _submit_attempt(attempt, request.POST)
        return redirect('exam_result', pk=attempt.pk)

    context = {
        'attempt': attempt,
        'exam': exam,
        'questions': questions,
        'existing_answers': existing_answers,
        'time_remaining': time_remaining,
        'total_questions': len(questions),
    }
    return render(request, 'exams/exam_take.html', context)


@login_required
@require_POST
def exam_submit(request, pk):
    attempt = get_object_or_404(ExamAttempt, pk=pk, student=request.user)
    if attempt.is_submitted:
        return redirect('exam_result', pk=attempt.pk)
    _submit_attempt(attempt, request.POST)
    return redirect('exam_result', pk=attempt.pk)


def _submit_attempt(attempt, post_data):
    """Process and save all answers, then auto-grade."""
    exam = attempt.exam
    # Only process and grade questions that were actually assigned to this specific attempt
    assigned_q_ids = attempt.answers.values_list('question_id', flat=True)
    questions = Question.objects.filter(id__in=assigned_q_ids).prefetch_related('options', 'blanks', 'pairs', 'order_items')

    for question in questions:
        answer, _ = AttemptAnswer.objects.get_or_create(attempt=attempt, question=question)
        qtype = question.question_type

        if qtype == 'mcq_single':
            selected = post_data.get(f'q_{question.id}', '')
            answer.selected_options = selected
            answer.answer_text = ''

        elif qtype == 'mcq_multi':
            if hasattr(post_data, 'getlist'):
                selected = post_data.getlist(f'q_{question.id}')
            else:
                val = post_data.get(f'q_{question.id}', [])
                selected = val if isinstance(val, list) else ([val] if val else [])
            answer.selected_options = ','.join(selected)
            answer.answer_text = ''

        elif qtype == 'true_false':
            answer.answer_text = post_data.get(f'q_{question.id}', '')

        elif qtype == 'fill_blank':
            answer.answer_text = post_data.get(f'q_{question.id}', '').strip()

        elif qtype in ['short_answer', 'essay']:
            answer.answer_text = post_data.get(f'q_{question.id}', '').strip()
            answer.is_graded = False

        elif qtype == 'matching':
            pairs = {}
            for pair in question.pairs.all():
                chosen = post_data.get(f'q_{question.id}_left_{pair.id}', '')
                if chosen:
                    pairs[str(pair.id)] = chosen
            answer.matching_answer = json.dumps(pairs)

        elif qtype == 'ordering':
            order_str = post_data.get(f'q_{question.id}_order', '')
            answer.ordering_answer = order_str

        elif qtype == 'code':
            answer.answer_text = post_data.get(f'q_{question.id}', '').strip()
            answer.is_graded = False

        answer.save()

    # Save time spent
    elapsed = (timezone.now() - attempt.started_at).total_seconds()
    attempt.time_spent_seconds = int(elapsed)
    attempt.is_submitted = True
    attempt.submitted_at = timezone.now()
    attempt.save()

    # Auto-grade
    auto_grade_attempt(attempt)


@login_required
@require_POST
def save_answer_ajax(request, attempt_pk):
    """AJAX endpoint to save answer progress"""
    attempt = get_object_or_404(ExamAttempt, pk=attempt_pk, student=request.user)
    if attempt.is_submitted:
        return JsonResponse({'status': 'submitted'})

    try:
        data = json.loads(request.body)
        question_id = data.get('question_id')
        question = get_object_or_404(Question, pk=question_id, exam=attempt.exam)
        answer, _ = AttemptAnswer.objects.get_or_create(attempt=attempt, question=question)

        qtype = question.question_type
        if qtype in ['mcq_single']:
            answer.selected_options = str(data.get('value', ''))
        elif qtype == 'mcq_multi':
            val = data.get('value', [])
            answer.selected_options = ','.join(str(v) for v in val)
        elif qtype in ['true_false', 'fill_blank', 'short_answer', 'essay']:
            answer.answer_text = str(data.get('value', ''))
        elif qtype == 'matching':
            answer.matching_answer = json.dumps(data.get('value', {}))
        elif qtype == 'ordering':
            val = data.get('value', [])
            answer.ordering_answer = ','.join(str(v) for v in val)

        answer.save()
        return JsonResponse({'status': 'ok'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})


@login_required
@require_POST
def log_violation_ajax(request, attempt_pk):
    """AJAX endpoint to log violations"""
    attempt = get_object_or_404(ExamAttempt, pk=attempt_pk, student=request.user)
    if attempt.is_submitted:
        return JsonResponse({'status': 'submitted'})

    try:
        data = json.loads(request.body)
        ViolationLog.objects.create(
            attempt=attempt,
            violation_type=data.get('type', 'tab_switch'),
            details=data.get('details', '')[:255],
        )
        attempt.violations_count += 1
        attempt.save(update_fields=['violations_count'])
        return JsonResponse({'status': 'ok', 'total': attempt.violations_count})
    except Exception as e:
        return JsonResponse({'status': 'error'})


# ─── RESULTS ──────────────────────────────────────────────────────────────────

@login_required
def exam_result(request, pk):
    attempt = get_object_or_404(ExamAttempt, pk=pk)
    # Access control
    try:
        profile = request.user.profile
        if profile.is_student() and attempt.student != request.user:
            return HttpResponseForbidden()
    except Profile.DoesNotExist:
        pass

    answers = attempt.answers.select_related('question').prefetch_related(
        'question__options', 'question__blanks', 'question__pairs', 'question__order_items'
    )

    context = {
        'attempt': attempt,
        'exam': attempt.exam,
        'answers': answers,
        'percentage': attempt.get_percentage(),
        'passed': attempt.get_percentage() >= attempt.exam.pass_mark / attempt.exam.total_marks * 100 if attempt.exam.total_marks else False,
    }
    return render(request, 'exams/exam_result.html', context)


@login_required
def results_list(request):
    try:
        profile = request.user.profile
    except Profile.DoesNotExist:
        return HttpResponseForbidden()

    exam_id = request.GET.get('exam')
    status_filter = request.GET.get('status', 'all')

    if profile.is_admin() or profile.is_assistant():
        if profile.is_admin() or request.user.is_superuser:
            attempts = ExamAttempt.objects.filter(
                exam__created_by=request.user, is_submitted=True
            ).select_related('student', 'exam').order_by('-submitted_at')
            exams = Exam.objects.filter(created_by=request.user)
        else:
            # Assistant sees all attempts for grading
            attempts = ExamAttempt.objects.filter(is_submitted=True).select_related('student', 'exam').order_by('-submitted_at')
            exams = Exam.objects.all()
    else:
        attempts = ExamAttempt.objects.filter(
            student=request.user, is_submitted=True
        ).select_related('exam').order_by('-submitted_at')
        exams = Exam.objects.filter(attempts__student=request.user).distinct()

    if exam_id:
        attempts = attempts.filter(exam_id=exam_id)
    if status_filter == 'pending':
        attempts = attempts.filter(is_fully_graded=False)
    elif status_filter == 'graded':
        attempts = attempts.filter(is_fully_graded=True)

    context = {
        'attempts': attempts,
        'exams': exams,
        'selected_exam': exam_id,
        'status_filter': status_filter,
        'profile': profile,
    }
    return render(request, 'exams/results_list.html', context)


@require_role(['admin', 'assistant'])
def grade_attempt(request, pk):
    profile = request.user.profile
    if profile.is_admin() or request.user.is_superuser:
        attempt = get_object_or_404(ExamAttempt, pk=pk, exam__created_by=request.user)
    else:
        # Assistant can grade any attempt
        attempt = get_object_or_404(ExamAttempt, pk=pk)
    answers = attempt.answers.select_related('question').prefetch_related(
        'question__options', 'question__blanks', 'question__pairs', 'question__order_items'
    ).order_by('question__order')

    if request.method == 'POST':
        for answer in answers:
            key_marks = f'marks_{answer.id}'
            key_comment = f'comment_{answer.id}'
            if key_marks in request.POST:
                try:
                    earned = float(request.POST[key_marks])
                    earned = max(0.0, min(earned, float(answer.question.marks)))
                    answer.earned_marks = earned
                    answer.grader_comment = request.POST.get(key_comment, '').strip()
                    answer.is_graded = True
                    answer.graded_at = timezone.now()
                    answer.graded_by = request.user
                    answer.save()
                except (ValueError, TypeError):
                    pass

        attempt.grader_notes = request.POST.get('grader_notes', '').strip()

        # Recalculate score
        total = sum(a.earned_marks for a in attempt.answers.filter(is_graded=True) if a.earned_marks is not None)
        attempt.final_score = total
        attempt.grade = attempt.calculate_grade()
        attempt.is_fully_graded = not attempt.answers.filter(is_graded=False).exists()
        attempt.save()

        messages.success(request, f'تم حفظ التصحيح. الدرجة النهائية: {total}/{attempt.exam.total_marks}')
        # Notify student that grading is done
        if attempt.is_fully_graded:
            create_notification(
                user=attempt.student,
                notification_type='grading_done',
                title=f'تم تصحيح: {attempt.exam.title}',
                message=f'درجتك النهائية: {total}/{attempt.exam.total_marks}',
                link=f'/attempts/{attempt.pk}/result/',
            )
        return redirect('results_list')

    context = {
        'attempt': attempt,
        'exam': attempt.exam,
        'answers': answers,
    }
    return render(request, 'exams/grade_attempt.html', context)


# ─── MONITORING ───────────────────────────────────────────────────────────────

@require_role(['admin', 'assistant'])
def monitoring(request):
    profile = request.user.profile
    # Superusers, Assistants, and Admins can see all active attempts for monitoring
    if request.user.is_superuser or profile.role in ['admin', 'assistant']:
        active_attempts = ExamAttempt.objects.filter(
            is_submitted=False,
        ).select_related('student', 'exam').order_by('-started_at')
    else:
        # Others see only what they created
        active_attempts = ExamAttempt.objects.filter(
            exam__created_by=request.user,
            is_submitted=False,
        ).select_related('student', 'exam').order_by('-started_at')

    # Annotate with time remaining
    for a in active_attempts:
        elapsed = (timezone.now() - a.started_at).total_seconds()
        a.time_remaining = max(0, int(a.exam.duration * 60 - elapsed))
        a.progress = min(100, int((a.answers.count() / max(1, a.exam.questions.count())) * 100))

    context = {
        'active_attempts': active_attempts,
        'exams': Exam.objects.filter(created_by=request.user, status='active') if profile.is_admin() else Exam.objects.filter(status='active'),
    }
    return render(request, 'exams/monitoring.html', context)


@require_role(['admin', 'assistant'])
def force_submit(request, attempt_pk):
    """Admin can force-submit an ongoing exam"""
    if request.method == 'POST':
        profile = request.user.profile
        if profile.is_admin() or request.user.is_superuser:
            attempt = get_object_or_404(ExamAttempt, pk=attempt_pk, exam__created_by=request.user)
        else:
            attempt = get_object_or_404(ExamAttempt, pk=attempt_pk)
        if not attempt.is_submitted:
            _submit_attempt(attempt, {})
            messages.warning(request, f'تم إنهاء جلسة {attempt.student.get_full_name()} قسراً')
    return redirect('monitoring')


# ─── API: Monitoring refresh ──────────────────────────────────────────────────

@require_role(['admin', 'assistant'])
def monitoring_data_api(request):
    """Returns JSON data for live monitoring refresh"""
    profile = request.user.profile
    if request.user.is_superuser or profile.role in ['admin', 'assistant']:
        active = ExamAttempt.objects.filter(is_submitted=False).select_related('student', 'exam')
    else:
        active = ExamAttempt.objects.filter(
            exam__created_by=request.user, is_submitted=False
        ).select_related('student', 'exam')

    data = []
    for a in active:
        elapsed = (timezone.now() - a.started_at).total_seconds()
        time_rem = max(0, int(a.exam.duration * 60 - elapsed))
        data.append({
            'id': a.id,
            'student': a.student.get_full_name(),
            'exam': a.exam.title,
            'time_remaining': time_rem,
            'violations': a.violations_count,
            'progress': min(100, int((a.answers.count() / max(1, a.exam.questions.count())) * 100)),
        })
    return JsonResponse({'attempts': data})


@require_role(['admin', 'assistant'])
def attempt_live_answers_api(request, attempt_pk):
    """API for the monitor to see a student's answers live."""
    attempt = get_object_or_404(ExamAttempt, pk=attempt_pk)
    
    # Permission check: monitor must be admin or the creator of the exam
    if not request.user.profile.is_admin() and attempt.exam.created_by != request.user:
        return JsonResponse({'error': 'Unauthorized'}, status=403)

    answers_map = {a.question_id: a for a in attempt.answers.all()}
    questions_data = []
    
    # Get questions in the order they appear in the attempt if possible, else default
    for q in attempt.exam.questions.all().order_by('id'):
        ans = answers_map.get(q.id)
        current_val = ""
        if ans:
            if q.question_type in ['mcq_single', 'true_false']:
                current_val = ans.selected_options if q.question_type == 'mcq_single' else ans.answer_text
                # For MCQ Single, let's find the option text
                if q.question_type == 'mcq_single' and current_val:
                    opt = q.options.filter(id=current_val).first()
                    if opt: current_val = opt.text
            elif q.question_type == 'mcq_multi':
                selected_ids = ans.selected_options.split(',') if ans.selected_options else []
                opts = q.options.filter(id__in=[x for x in selected_ids if x.strip()])
                current_val = ", ".join([o.text for o in opts])
            elif q.question_type in ['short_answer', 'essay', 'fill_blank']:
                current_val = ans.decrypted_text
            elif q.question_type == 'matching':
                current_val = "إجابة مطابقة مسجلة" # Simplified for monitor
            elif q.question_type == 'ordering':
                current_val = "ترتيب مسجل"
            elif q.question_type == 'code':
                current_val = ans.decrypted_text
        
        questions_data.append({
            'id': q.id,
            'text': q.text,
            'type': q.get_question_type_display(),
            'answer': current_val or "لم تتم الإجابة بعد",
            'is_answered': ans is not None
        })
        
    return JsonResponse({
        'student': attempt.student.get_full_name(),
        'exam': attempt.exam.title,
        'questions': questions_data
    })


# ─── STUDENTS LIST ────────────────────────────────────────────────────────────

@require_role(['admin', 'assistant'])
def students_list(request):
    from accounts.models import College, Department
    
    students = User.objects.filter(profile__role='student').select_related(
        'profile', 'profile__college', 'profile__academic_department',
        'profile__academic_department__college'
    ).order_by('last_name', 'first_name')
    
    # Filters
    college_id = request.GET.get('college', '')
    dept_id = request.GET.get('department', '')
    level_filter = request.GET.get('level', '')
    search_q = request.GET.get('q', '').strip()
    
    if college_id:
        students = students.filter(
            Q(profile__academic_department__college_id=college_id) |
            Q(profile__college_id=college_id)
        )
    if dept_id:
        students = students.filter(profile__academic_department_id=dept_id)
    if level_filter:
        students = students.filter(profile__level=level_filter)
    if search_q:
        students = students.filter(
            Q(first_name__icontains=search_q) |
            Q(last_name__icontains=search_q) |
            Q(username__icontains=search_q) |
            Q(profile__student_id__icontains=search_q)
        )
    
    colleges = College.objects.all()
    departments = Department.objects.all()
    if college_id:
        departments = departments.filter(college_id=college_id)
    
    # Stats
    total_all = User.objects.filter(profile__role='student').count()
    
    context = {
        'students': students,
        'total': students.count(),
        'total_all': total_all,
        'colleges': colleges,
        'departments': departments,
        'selected_college': int(college_id) if college_id else None,
        'selected_dept': int(dept_id) if dept_id else None,
        'selected_level': int(level_filter) if level_filter else None,
        'search_q': search_q,
    }
    return render(request, 'exams/students_list.html', context)


# ─── EXAM STATISTICS ─────────────────────────────────────────────────────────

@require_role(['admin', 'assistant'])
def exam_statistics(request, pk):
    exam = get_object_or_404(Exam, pk=pk, created_by=request.user)
    attempts = exam.attempts.filter(is_submitted=True).select_related('student')
    total_attempts = attempts.count()

    if total_attempts == 0:
        context = {
            'exam': exam,
            'no_data': True,
        }
        return render(request, 'exams/exam_statistics.html', context)

    scores = [a.get_total_score() for a in attempts]
    percentages = [a.get_percentage() for a in attempts]

    avg_score = round(sum(scores) / len(scores), 1)
    max_score = max(scores)
    min_score = min(scores)
    pass_threshold = (exam.pass_mark / exam.total_marks * 100) if exam.total_marks else 50
    passed_count = sum(1 for p in percentages if p >= pass_threshold)
    failed_count = total_attempts - passed_count

    # Grade distribution
    grade_dist = {}
    for a in attempts:
        g = a.grade or a.calculate_grade()
        grade_dist[g] = grade_dist.get(g, 0) + 1

    # Score distribution for chart (buckets of 10%)
    score_buckets = [0] * 10
    for p in percentages:
        bucket = min(int(p // 10), 9)
        score_buckets[bucket] += 1

    # Question difficulty analysis
    questions = exam.questions.all()
    question_stats = []
    for q in questions:
        q_answers = AttemptAnswer.objects.filter(
            question=q, attempt__is_submitted=True, is_graded=True
        )
        total_q = q_answers.count()
        if total_q > 0:
            correct_q = q_answers.filter(earned_marks=q.marks).count()
            difficulty = round((1 - correct_q / total_q) * 100, 1)
        else:
            correct_q = 0
            difficulty = 0
        question_stats.append({
            'question': q,
            'total': total_q,
            'correct': correct_q,
            'difficulty': difficulty,
        })

    context = {
        'exam': exam,
        'total_attempts': total_attempts,
        'avg_score': avg_score,
        'max_score': max_score,
        'min_score': min_score,
        'passed_count': passed_count,
        'failed_count': failed_count,
        'pass_rate': round(passed_count / total_attempts * 100, 1),
        'grade_dist': json.dumps(grade_dist),
        'grade_labels': json.dumps(list(grade_dist.keys())),
        'grade_values': json.dumps(list(grade_dist.values())),
        'score_buckets': json.dumps(score_buckets),
        'question_stats': question_stats,
        'no_data': False,
    }
    return render(request, 'exams/exam_statistics.html', context)


# ─── EXAM DUPLICATE ──────────────────────────────────────────────────────────

@require_role('admin')
def exam_duplicate(request, pk):
    original = get_object_or_404(Exam, pk=pk, created_by=request.user)
    if request.method == 'POST':
        new_exam = Exam.objects.create(
            title=f"{original.title} (نسخة)",
            subject=original.subject,
            description=original.description,
            duration=original.duration,
            total_marks=original.total_marks,
            pass_mark=original.pass_mark,
            max_attempts=original.max_attempts,
            shuffle_questions=original.shuffle_questions,
            allow_review=original.allow_review,
            status='draft',
            created_by=request.user,
        )
        # Duplicate questions
        for q in original.questions.all():
            new_q = Question.objects.create(
                exam=new_exam,
                question_type=q.question_type,
                text=q.text,
                marks=q.marks,
                order=q.order,
                explanation=q.explanation,
                tf_answer=q.tf_answer,
                rubric=q.rubric,
            )
            for opt in q.options.all():
                QuestionOption.objects.create(
                    question=new_q, text=opt.text,
                    is_correct=opt.is_correct, order=opt.order,
                )
            for blank in q.blanks.all():
                QuestionBlank.objects.create(
                    question=new_q, accepted_answer=blank.accepted_answer,
                    case_sensitive=blank.case_sensitive,
                )
            for pair in q.pairs.all():
                MatchingPair.objects.create(
                    question=new_q, left_text=pair.left_text,
                    right_text=pair.right_text, order=pair.order,
                )
            for item in q.order_items.all():
                OrderingItem.objects.create(
                    question=new_q, text=item.text,
                    correct_position=item.correct_position,
                )

        messages.success(request, f'تم نسخ الاختبار بنجاح إلى "{new_exam.title}"')
        return redirect('exam_edit', pk=new_exam.pk)
    return redirect('exam_list')


# ─── EXPORT RESULTS TO CSV ─────────────────────────────────────────────────────

@require_role(['admin', 'assistant'])
def export_results(request):
    exam_id = request.GET.get('exam')
    profile = request.user.profile
    if profile.is_admin() or request.user.is_superuser:
        attempts = ExamAttempt.objects.filter(
            exam__created_by=request.user, is_submitted=True
        ).select_related('student', 'exam').order_by('exam__title', '-submitted_at')
    else:
        attempts = ExamAttempt.objects.filter(
            is_submitted=True
        ).select_related('student', 'exam').order_by('exam__title', '-submitted_at')

    if exam_id:
        attempts = attempts.filter(exam_id=exam_id)

    response = HttpResponse(content_type='text/csv; charset=utf-8-sig')
    response['Content-Disposition'] = 'attachment; filename="exam_results.csv"'
    response.write('\ufeff')  # BOM for Excel Arabic support

    writer = csv.writer(response)
    writer.writerow(['الاختبار', 'اسم الطالب', 'اسم المستخدم', 'الدرجة', 'الدرجة الكلية', 'النسبة', 'التقدير', 'المخالفات', 'الحالة', 'تاريخ التسليم'])

    for a in attempts:
        writer.writerow([
            a.exam.title,
            a.student.get_full_name(),
            a.student.username,
            a.final_score or 0,
            a.exam.total_marks,
            f"{a.get_percentage()}%",
            a.grade or '-',
            a.violations_count,
            'مصحح' if a.is_fully_graded else 'بانتظار التصحيح',
            a.submitted_at.strftime('%Y-%m-%d %H:%M') if a.submitted_at else '',
        ])

    return response


# ─── NOTIFICATIONS ────────────────────────────────────────────────────────────

@login_required
def notifications_api(request):
    notifications = Notification.objects.filter(user=request.user)[:20]
    unread_count = Notification.objects.filter(user=request.user, is_read=False).count()
    data = {
        'unread_count': unread_count,
        'notifications': [
            {
                'id': n.id,
                'type': n.notification_type,
                'title': n.title,
                'message': n.message,
                'link': n.link,
                'is_read': n.is_read,
                'created_at': n.created_at.strftime('%Y-%m-%d %H:%M'),
            }
            for n in notifications
        ]
    }
    return JsonResponse(data)


@login_required
@require_POST
def mark_notifications_read(request):
    Notification.objects.filter(user=request.user, is_read=False).update(is_read=True)
    return JsonResponse({'status': 'ok'})

@login_required
@require_POST
def contact_api_submit(request):
    try:
        import json
        data = json.loads(request.body)
        subject = data.get('subject')
        message = data.get('message')
        
        if not subject or not message:
            return JsonResponse({'status': 'error', 'msg': 'يرجى إكمال جميع الحقول'})
            
        from .models import ContactMessage
        ContactMessage.objects.create(
            user=request.user,
            subject=subject,
            message=message
        )
        return JsonResponse({'status': 'success', 'msg': 'تم استلام رسالتك، سيصلك الرد قريباً عبر الإشعارات.'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'msg': str(e)})


def create_notification(user, notification_type, title, message='', link=''):
    """Helper to create a notification for a user."""
    Notification.objects.create(
        user=user,
        notification_type=notification_type,
        title=title,
        message=message,
        link=link,
    )
@login_required
def generate_certificate(request, attempt_pk):
    """
    Renders a premium, printable certificate for students who passed.
    """
    # Allow Admins/Superusers to see ANY certificate, but students only THEIR OWN
    profile = getattr(request.user, 'profile', None)
    is_staff = (profile and profile.role in ['admin', 'assistant']) or request.user.is_superuser
    
    if is_staff:
        attempt = get_object_or_404(ExamAttempt, pk=attempt_pk, is_submitted=True)
    else:
        attempt = get_object_or_404(ExamAttempt, pk=attempt_pk, student=request.user, is_submitted=True)
    
    # Check if passed and fully graded
    if not attempt.exam.total_marks or attempt.exam.total_marks == 0:
        return HttpResponseForbidden("لا يمكن إصدار شهادة لاختبار بدون درجات.")
        
    percentage = (attempt.get_total_score() / attempt.exam.total_marks) * 100
    passed = percentage >= attempt.exam.pass_mark
    
    if not passed:
        return HttpResponseForbidden("الشهادة متاحة فقط للاختبارات المجتازة.")
    
    if not attempt.is_fully_graded:
        return HttpResponseForbidden("الشهادة ستكون متاحة فور اكتمال تصحيح جميع الأسئلة.")
        
    context = {
        'attempt': attempt,
        'percentage': round(percentage, 1),
        'passed': passed,
        'now': timezone.now(),
    }
    return render(request, 'exams/certificate.html', context)


@login_required
@require_POST
def send_chat_message(request, attempt_id):
    attempt = get_object_or_404(ExamAttempt, pk=attempt_id)
    profile = getattr(request.user, 'profile', None)
    is_owner = attempt.student == request.user
    is_staff = (profile and profile.role in ['admin', 'assistant']) or request.user.is_superuser
    
    if not (is_owner or is_staff):
        return HttpResponseForbidden("غير مسموح لك بإرسال رسائل هنا.")
        
    try:
        data = json.loads(request.body)
        msg_text = data.get('message', '').strip()
    except:
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

    if msg_text:
        msg = ChatMessage.objects.create(
            attempt=attempt,
            sender=request.user,
            message=msg_text
        )
        return JsonResponse({
            'status': 'ok',
            'id': msg.id,
            'sender': 'أنت' if msg.sender == request.user else msg.sender.username,
            'timestamp': msg.timestamp.strftime('%H:%M')
        })
    return JsonResponse({'status': 'error', 'message': 'رسالة فارغة'}, status=400)


@login_required
def get_chat_messages(request, attempt_id):
    attempt = get_object_or_404(ExamAttempt, pk=attempt_id)
    profile = getattr(request.user, 'profile', None)
    is_owner = attempt.student == request.user
    is_staff = (profile and profile.role in ['admin', 'assistant']) or request.user.is_superuser
    
    if not (is_owner or is_staff):
        return HttpResponseForbidden()
        
    messages = ChatMessage.objects.filter(attempt=attempt).order_by('timestamp')
    return JsonResponse({
        'messages': [
            {
                'id': m.id,
                'sender': 'أنت' if m.sender == request.user else m.sender.username,
                'is_me': m.sender == request.user,
                'message': m.message,
                'timestamp': m.timestamp.strftime('%H:%M')
            }
            for m in messages
        ]
    })

@require_role(['admin', 'assistant'])
def import_questions(request, exam_pk):
    exam = get_object_or_404(Exam, pk=exam_pk)
    
    # Permission: Admins import for everything, others only their own
    if not request.user.profile.is_admin() and exam.created_by != request.user:
        return HttpResponseForbidden("غير مصرح لك باستيراد أسئلة لهذا الاختبار")
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']
        ext = file.name.split('.')[-1].lower()
        
        count = 0
        try:
            if ext == 'xlsx':
                wb = openpyxl.load_workbook(file)
                sheet = wb.active
                for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True)):
                    if not row[0]: continue  # Text is required
                    q = Question.objects.create(
                        exam=exam,
                        text=row[0],
                        question_type=row[1] if row[1] else 'mcq_single',
                        marks=int(row[2]) if row[2] else 5,
                        explanation=row[3] if row[3] else '',
                    )
                    # Handle options if MCQ
                    if q.question_type in ['mcq_single', 'mcq_multi']:
                        options_str = str(row[4]) if row[4] else ''
                        correct_indices = str(row[5]).split(',') if row[5] is not None else []
                        for i, opt_text in enumerate(options_str.split('|')):
                            if opt_text.strip():
                                QuestionOption.objects.create(
                                    question=q, text=opt_text.strip(),
                                    is_correct=str(i) in correct_indices, order=i
                                )
                    elif q.question_type == 'true_false':
                        q.tf_answer = str(row[4]).lower() in ['true', 'yes', '1', 'صح']
                        q.save()
                    count += 1
                    
            elif ext == 'docx':
                doc = Document(file)
                # Basic parsing for Word: Question: ... \n Type: ... \n Marks: ... \n - Option (Correct)
                current_q = None
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text.startswith('Question:'):
                        current_q = Question.objects.create(
                            exam=exam, text=text.replace('Question:', '').strip(),
                            question_type='mcq_single', marks=5
                        )
                        count += 1
                    elif current_q and text.startswith('Type:'):
                        current_q.question_type = text.replace('Type:', '').strip()
                        current_q.save()
                    elif current_q and text.startswith('Marks:'):
                        try: current_q.marks = int(text.replace('Marks:', '').strip())
                        except: pass
                        current_q.save()
                    elif current_q and text.startswith('-'):
                        is_correct = '(Correct)' in text
                        opt_text = text.replace('-', '').replace('(Correct)', '').strip()
                        QuestionOption.objects.create(
                            question=current_q, text=opt_text, is_correct=is_correct
                        )
            
            messages.success(request, f'تم استيراد {count} سؤال بنجاح')
        except Exception as e:
            messages.error(request, f'خطأ أثناء الاستيراد: {str(e)}')
            
        return redirect('exam_edit', pk=exam_pk)
    return redirect('exam_edit', pk=exam_pk)


@login_required
@require_POST
def upload_proctor_snapshot(request, attempt_pk):
    attempt = get_object_or_404(ExamAttempt, pk=attempt_pk, student=request.user)
    if attempt.is_submitted:
        return JsonResponse({'status': 'error', 'message': 'Attempt already submitted'})
        
    if 'image' in request.FILES:
        snapshot = ProctorSnapshot.objects.create(
            attempt=attempt,
            image=request.FILES['image']
        )
        return JsonResponse({'status': 'ok', 'id': snapshot.id})
    return JsonResponse({'status': 'error', 'message': 'No image provided'}, status=400)


@require_role(['admin', 'assistant'])
def attempt_snapshots_api(request, attempt_pk):
    attempt = get_object_or_404(ExamAttempt, pk=attempt_pk)
    # Access control
    profile = getattr(request.user, 'profile', None)
    is_staff = (profile and profile.role in ['admin', 'assistant']) or request.user.is_superuser
    if not is_staff and attempt.student != request.user:
        return HttpResponseForbidden()
        
    snapshots = attempt.snapshots.all()
    data = [{
        'id': s.id,
        'url': s.image.url,
        'timestamp': s.timestamp.strftime('%H:%M:%S')
    } for s in snapshots]
    return JsonResponse({'snapshots': data})


def about_system(request):
    """Renders the About page with FAQ and contact info."""
    return render(request, 'exams/about.html')

def download_sample_excel(request):
    """Generates a sample Excel file for bulk question import."""
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.title = "Sample Questions"
    headers = ["Text", "Type", "Marks", "Explanation", "Options/Answer", "Correct Indices"]
    sheet.append(headers)
    sheet.append(["ما هي عاصمة العراق؟", "mcq_single", 5, "بغداد هي العاصمة السياسية.", "بغداد|البصرة|الموصل|أربيل", "0"])
    sheet.append(["اختر لغات البرمجة (متعدد):", "mcq_multi", 10, "", "Python|Java|HTML|CSS", "0,1"])
    sheet.append(["الأرض مسطحة؟", "true_false", 5, "الأرض كروية.", "خطأ", ""])
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=sample_questions.xlsx'
    wb.save(response)
    return response

def download_sample_word(request):
    """Generates a sample Word file for bulk question import."""
    doc = Document()
    doc.add_heading('نموذج استيراد الأسئلة', 0)
    doc.add_paragraph('Question: ما هو ناتج 5 + 5؟')
    doc.add_paragraph('Type: mcq_single')
    doc.add_paragraph('Marks: 5')
    doc.add_paragraph('- 10 (Correct)')
    doc.add_paragraph('- 15')
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=sample_questions.docx'
    doc.save(response)
    return response

@require_role(['admin', 'assistant'])
def export_questions(request, exam_pk):
    exam = get_object_or_404(Exam, pk=exam_pk)
    
    # Permission: Admins export everything, others only their own
    if not request.user.profile.is_admin() and exam.created_by != request.user:
        return HttpResponseForbidden("غير مصرح لك بتصدير أسئلة هذا الاختبار")
        
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.title = "Questions"
    
    # Headers matching the import format
    headers = ["Text", "Type", "Marks", "Explanation", "Options/Answer", "Correct Indices"]
    sheet.append(headers)
    
    for q in exam.questions.all():
        row = [
            q.text,
            q.question_type,
            q.marks,
            q.explanation or ""
        ]
        
        # Format options or answer
        if q.question_type in ['mcq_single', 'mcq_multi']:
            options = q.options.all().order_by('order')
            row.append("|".join([opt.text for opt in options]))
            row.append(",".join([str(i) for i, opt in enumerate(options) if opt.is_correct]))
        elif q.question_type == 'true_false':
            row.append("صح" if q.tf_answer else "خطأ")
            row.append("")
        else:
            row.append("")
            row.append("")
            
        sheet.append(row)
        
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename=questions_{exam_pk}.xlsx'
    wb.save(response)
    return response


# ═══════════════════════════════════════════════════════════════════════════════
# ─── ACCOUNT MANAGEMENT SYSTEM ────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

import string
import uuid

def _generate_password(length=8):
    """Generate a random password."""
    chars = string.ascii_letters + string.digits
    return ''.join(random.choice(chars) for _ in range(length))

def _generate_username(first_name, last_name, student_id=''):
    """Generate a unique username."""
    if student_id:
        base = student_id.strip()
    else:
        base = f"{first_name.strip()}.{last_name.strip()}".lower()
        # Transliterate Arabic
        base = base.replace(' ', '_')
    
    username = base
    counter = 1
    while User.objects.filter(username=username).exists():
        username = f"{base}{counter}"
        counter += 1
    return username


@require_role('admin')
def account_management(request):
    """Main account management page."""
    from accounts.models import College, Department
    
    colleges = College.objects.all()
    departments = Department.objects.all()
    
    # Filters
    role_filter = request.GET.get('role', '')
    college_filter = request.GET.get('college', '')
    dept_filter = request.GET.get('department', '')
    batch_filter = request.GET.get('batch', '')
    search_q = request.GET.get('q', '').strip()
    
    # Get generated accounts
    generated = Profile.objects.filter(created_by_admin=True).select_related(
        'user', 'college', 'academic_department'
    ).order_by('-user__date_joined')
    
    if role_filter:
        generated = generated.filter(role=role_filter)
    if college_filter:
        generated = generated.filter(
            Q(college_id=college_filter) | Q(academic_department__college_id=college_filter)
        )
    if dept_filter:
        generated = generated.filter(academic_department_id=dept_filter)
    if batch_filter:
        generated = generated.filter(batch_id=batch_filter)
    if search_q:
        generated = generated.filter(
            Q(user__first_name__icontains=search_q) |
            Q(user__last_name__icontains=search_q) |
            Q(user__username__icontains=search_q) |
            Q(student_id__icontains=search_q)
        )
    
    # Get unique batches for filter
    batches = Profile.objects.filter(
        created_by_admin=True, batch_id__gt=''
    ).values_list('batch_id', flat=True).distinct().order_by('-batch_id')[:20]
    
    # Stats
    total_generated = Profile.objects.filter(created_by_admin=True).count()
    total_students_gen = Profile.objects.filter(created_by_admin=True, role='student').count()
    total_admins_gen = Profile.objects.filter(created_by_admin=True, role='admin').count()
    total_assistants_gen = Profile.objects.filter(created_by_admin=True, role='assistant').count()
    
    context = {
        'colleges': colleges,
        'departments': departments,
        'generated_accounts': generated,
        'batches': batches,
        'role_filter': role_filter,
        'college_filter': int(college_filter) if college_filter else None,
        'dept_filter': int(dept_filter) if dept_filter else None,
        'batch_filter': batch_filter,
        'search_q': search_q,
        'total_generated': total_generated,
        'total_students_gen': total_students_gen,
        'total_admins_gen': total_admins_gen,
        'total_assistants_gen': total_assistants_gen,
    }
    return render(request, 'exams/account_management.html', context)


@require_role('admin')
@require_POST
def generate_single_account(request):
    """Generate a single student/instructor account."""
    from accounts.models import College, Department
    
    first_name = request.POST.get('first_name', '').strip()
    last_name = request.POST.get('last_name', '').strip()
    email = request.POST.get('email', '').strip()
    role = request.POST.get('role', 'student')
    college_id = request.POST.get('college')
    dept_id = request.POST.get('department')
    student_id = request.POST.get('student_id', '').strip()
    level = request.POST.get('level')
    phone = request.POST.get('phone', '').strip()
    
    if not first_name or not last_name:
        messages.error(request, 'يرجى إدخال الاسم الأول والأخير')
        return redirect('account_management')
    
    username = _generate_username(first_name, last_name, student_id)
    password = _generate_password()
    
    user = User.objects.create_user(
        username=username,
        email=email,
        password=password,
        first_name=first_name,
        last_name=last_name,
    )
    
    if role == 'admin':
        user.is_staff = True
        user.save()
    
    profile, _ = Profile.objects.get_or_create(user=user)
    profile.role = role
    profile.student_id = student_id
    profile.college_id = college_id if college_id else None
    profile.academic_department_id = dept_id if dept_id else None
    profile.level = int(level) if level else None
    profile.phone = phone
    profile.generated_password = password
    profile.created_by_admin = True
    profile.batch_id = f"single_{timezone.now().strftime('%Y%m%d_%H%M%S')}"
    profile.save()
    
    messages.success(request, f'تم إنشاء الحساب بنجاح: {username} / {password}')
    return redirect('account_management')


@require_role('admin')
@require_POST
def generate_bulk_accounts(request):
    """Generate multiple accounts from an Excel file."""
    from accounts.models import College, Department
    
    file = request.FILES.get('file')
    role = request.POST.get('role', 'student')
    college_id = request.POST.get('college')
    dept_id = request.POST.get('department')
    level = request.POST.get('level')
    
    if not file:
        messages.error(request, 'يرجى رفع ملف Excel')
        return redirect('account_management')
    
    batch_id = f"batch_{timezone.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}"
    count = 0
    errors = []
    
    try:
        wb = openpyxl.load_workbook(file)
        sheet = wb.active
        
        for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
            if not row or not row[0]:
                continue
            
            first_name = str(row[0]).strip()
            last_name = str(row[1]).strip() if len(row) > 1 and row[1] else ''
            email = str(row[2]).strip() if len(row) > 2 and row[2] else ''
            sid = str(row[3]).strip() if len(row) > 3 and row[3] else ''
            phone_val = str(row[4]).strip() if len(row) > 4 and row[4] else ''
            
            if not first_name:
                continue
            
            try:
                username = _generate_username(first_name, last_name, sid)
                password = _generate_password()
                
                user = User.objects.create_user(
                    username=username,
                    email=email,
                    password=password,
                    first_name=first_name,
                    last_name=last_name,
                )
                
                if role == 'admin':
                    user.is_staff = True
                    user.save()
                
                profile, _ = Profile.objects.get_or_create(user=user)
                profile.role = role
                profile.student_id = sid
                profile.college_id = college_id if college_id else None
                profile.academic_department_id = dept_id if dept_id else None
                profile.level = int(level) if level else None
                profile.phone = phone_val
                profile.generated_password = password
                profile.created_by_admin = True
                profile.batch_id = batch_id
                profile.save()
                count += 1
            except Exception as e:
                errors.append(f"صف {row_idx}: {str(e)}")
        
        if count > 0:
            messages.success(request, f'تم إنشاء {count} حساب بنجاح (دفعة: {batch_id})')
        if errors:
            messages.warning(request, f'حدثت {len(errors)} أخطاء أثناء الاستيراد')
            
    except Exception as e:
        messages.error(request, f'خطأ في قراءة الملف: {str(e)}')
    
    return redirect('account_management')


@require_role('admin')
def export_accounts_excel(request):
    """Export generated accounts to Excel."""
    batch = request.GET.get('batch', '')
    role_filter = request.GET.get('role', '')
    college_filter = request.GET.get('college', '')
    dept_filter = request.GET.get('department', '')
    
    accounts = Profile.objects.filter(created_by_admin=True).select_related(
        'user', 'college', 'academic_department', 'academic_department__college'
    )
    
    if batch:
        accounts = accounts.filter(batch_id=batch)
    if role_filter:
        accounts = accounts.filter(role=role_filter)
    if college_filter:
        accounts = accounts.filter(
            Q(college_id=college_filter) | Q(academic_department__college_id=college_filter)
        )
    if dept_filter:
        accounts = accounts.filter(academic_department_id=dept_filter)
    
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.title = "الحسابات المولّدة"
    
    headers = ['#', 'الاسم الكامل', 'اسم المستخدم', 'كلمة المرور', 'الدور',
               'الكلية', 'القسم', 'الرقم الجامعي', 'المستوى', 'البريد', 'الهاتف']
    sheet.append(headers)
    
    # Style header
    from openpyxl.styles import Font, Alignment, PatternFill
    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='1a1a2e', end_color='1a1a2e', fill_type='solid')
    for cell in sheet[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    
    for idx, acc in enumerate(accounts, 1):
        sheet.append([
            idx,
            acc.user.get_full_name(),
            acc.user.username,
            acc.generated_password or '***',
            acc.get_role_display(),
            acc.get_college_name(),
            acc.get_department_name(),
            acc.student_id or '-',
            acc.get_level_display_safe(),
            acc.user.email or '-',
            acc.phone or '-',
        ])
    
    # Auto-width columns
    for col in sheet.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        sheet.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="generated_accounts.xlsx"'
    response.write('\ufeff'.encode('utf-8'))
    wb.save(response)
    return response


@require_role('admin')
def print_account_cards(request):
    """Print account cards for generated accounts."""
    batch = request.GET.get('batch', '')
    ids = request.GET.get('ids', '')
    
    accounts = Profile.objects.filter(created_by_admin=True).select_related(
        'user', 'college', 'academic_department'
    )
    
    if batch:
        accounts = accounts.filter(batch_id=batch)
    elif ids:
        id_list = [int(x) for x in ids.split(',') if x.strip().isdigit()]
        accounts = accounts.filter(id__in=id_list)
    
    return render(request, 'exams/print_accounts.html', {
        'accounts': accounts,
        'batch': batch,
    })


@require_role('admin')
def download_bulk_template(request):
    """Download a sample Excel template for bulk account creation."""
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.title = "قالب الحسابات"
    
    headers = ['الاسم الأول', 'الاسم الأخير', 'البريد الإلكتروني', 'الرقم الجامعي', 'رقم الهاتف']
    sheet.append(headers)
    
    # Sample data
    sheet.append(['أحمد', 'محمد', 'ahmed@example.com', 'STD001', '0771234567'])
    sheet.append(['سارة', 'علي', 'sara@example.com', 'STD002', '0779876543'])
    sheet.append(['خالد', 'حسن', 'khaled@example.com', 'STD003', '0771112233'])
    
    from openpyxl.styles import Font, PatternFill, Alignment
    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='16213e', end_color='16213e', fill_type='solid')
    for cell in sheet[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    
    for col in sheet.columns:
        sheet.column_dimensions[col[0].column_letter].width = 25
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="bulk_accounts_template.xlsx"'
    wb.save(response)
    return response


# ═══════════════════════════════════════════════════════════════════════════════
# ─── ACADEMIC STRUCTURE MANAGEMENT ────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

@require_role('admin')
def academic_structure(request):
    """Manage colleges, departments, and courses."""
    from accounts.models import College, Department
    
    colleges = College.objects.prefetch_related('departments').all()
    courses = Course.objects.select_related('department', 'department__college', 'instructor').all()
    
    context = {
        'colleges': colleges,
        'courses': courses,
        'total_colleges': colleges.count(),
        'total_departments': Department.objects.count(),
        'total_courses': courses.count(),
    }
    return render(request, 'exams/academic_structure.html', context)


@require_role('admin')
@require_POST
def add_college(request):
    from accounts.models import College
    name = request.POST.get('name', '').strip()
    code = request.POST.get('code', '').strip()
    if name:
        College.objects.create(name=name, code=code)
        messages.success(request, f'تم إضافة الكلية "{name}" بنجاح')
    return redirect('academic_structure')


@require_role('admin')
@require_POST
def edit_college(request, pk):
    from accounts.models import College
    college = get_object_or_404(College, pk=pk)
    college.name = request.POST.get('name', college.name).strip()
    college.code = request.POST.get('code', college.code).strip()
    college.save()
    messages.success(request, f'تم تحديث الكلية "{college.name}"')
    return redirect('academic_structure')


@require_role('admin')
@require_POST
def delete_college(request, pk):
    from accounts.models import College
    college = get_object_or_404(College, pk=pk)
    name = college.name
    college.delete()
    messages.success(request, f'تم حذف الكلية "{name}"')
    return redirect('academic_structure')


@require_role('admin')
@require_POST
def add_department(request):
    from accounts.models import Department
    name = request.POST.get('name', '').strip()
    code = request.POST.get('code', '').strip()
    college_id = request.POST.get('college')
    if name and college_id:
        Department.objects.create(name=name, code=code, college_id=college_id)
        messages.success(request, f'تم إضافة القسم "{name}" بنجاح')
    return redirect('academic_structure')


@require_role('admin')
@require_POST
def edit_department(request, pk):
    from accounts.models import Department
    dept = get_object_or_404(Department, pk=pk)
    dept.name = request.POST.get('name', dept.name).strip()
    dept.code = request.POST.get('code', dept.code).strip()
    dept.college_id = request.POST.get('college', dept.college_id)
    dept.save()
    messages.success(request, f'تم تحديث القسم "{dept.name}"')
    return redirect('academic_structure')


@require_role('admin')
@require_POST
def delete_department(request, pk):
    from accounts.models import Department
    dept = get_object_or_404(Department, pk=pk)
    name = dept.name
    dept.delete()
    messages.success(request, f'تم حذف القسم "{name}"')
    return redirect('academic_structure')


@require_role('admin')
@require_POST
def add_course(request):
    name = request.POST.get('name', '').strip()
    code = request.POST.get('code', '').strip()
    dept_id = request.POST.get('department')
    instructor_id = request.POST.get('instructor')
    level = request.POST.get('level')
    semester = request.POST.get('semester', '')
    
    if name and dept_id:
        Course.objects.create(
            name=name, code=code, department_id=dept_id,
            instructor_id=instructor_id if instructor_id else None,
            level=int(level) if level else None,
            semester=semester,
        )
        messages.success(request, f'تم إضافة المادة "{name}" بنجاح')
    return redirect('academic_structure')


@require_role('admin')
@require_POST  
def edit_course(request, pk):
    course = get_object_or_404(Course, pk=pk)
    course.name = request.POST.get('name', course.name).strip()
    course.code = request.POST.get('code', course.code).strip()
    course.department_id = request.POST.get('department', course.department_id)
    instructor_id = request.POST.get('instructor')
    course.instructor_id = instructor_id if instructor_id else None
    level = request.POST.get('level')
    course.level = int(level) if level else None
    course.semester = request.POST.get('semester', course.semester)
    course.save()
    messages.success(request, f'تم تحديث المادة "{course.name}"')
    return redirect('academic_structure')


@require_role('admin')
@require_POST
def delete_course(request, pk):
    course = get_object_or_404(Course, pk=pk)
    name = course.name
    course.delete()
    messages.success(request, f'تم حذف المادة "{name}"')
    return redirect('academic_structure')


# ═══════════════════════════════════════════════════════════════════════════════
# ─── CASCADING API ENDPOINTS ─────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

@login_required
def api_departments_by_college(request):
    """Return departments for a given college (AJAX)."""
    from accounts.models import Department
    college_id = request.GET.get('college_id')
    if college_id:
        depts = Department.objects.filter(college_id=college_id).values('id', 'name', 'code')
    else:
        depts = Department.objects.all().values('id', 'name', 'code')
    return JsonResponse({'departments': list(depts)})


@login_required
def api_courses_by_department(request):
    """Return courses for a given department (AJAX)."""
    dept_id = request.GET.get('department_id')
    if dept_id:
        courses = Course.objects.filter(department_id=dept_id).values('id', 'name', 'code')
    else:
        courses = Course.objects.all().values('id', 'name', 'code')
    return JsonResponse({'courses': list(courses)})


@login_required
def api_students_by_department(request):
    """Return students for a given department (AJAX)."""
    dept_id = request.GET.get('department_id')
    college_id = request.GET.get('college_id')
    
    students = User.objects.filter(profile__role='student')
    if dept_id:
        students = students.filter(profile__academic_department_id=dept_id)
    elif college_id:
        students = students.filter(profile__academic_department__college_id=college_id)
    
    data = [{
        'id': s.id,
        'name': s.get_full_name() or s.username,
        'student_id': s.profile.student_id if hasattr(s, 'profile') else '',
    } for s in students.select_related('profile')[:200]]
    
    return JsonResponse({'students': data})


# ═══════════════════════════════════════════════════════════════════════════════
# ─── ENHANCED STUDENTS LIST ──────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

@require_role(['admin', 'assistant'])
def export_students_excel(request):
    """Export students list to Excel."""
    from accounts.models import College, Department
    
    students = User.objects.filter(profile__role='student').select_related('profile', 'profile__college', 'profile__academic_department')
    
    college_id = request.GET.get('college')
    dept_id = request.GET.get('department')
    if college_id:
        students = students.filter(profile__academic_department__college_id=college_id)
    if dept_id:
        students = students.filter(profile__academic_department_id=dept_id)
    
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.title = "قائمة الطلاب"
    
    headers = ['#', 'الاسم الكامل', 'اسم المستخدم', 'الرقم الجامعي', 'الكلية', 'القسم', 'المستوى', 'البريد', 'الهاتف']
    sheet.append(headers)
    
    from openpyxl.styles import Font, PatternFill, Alignment
    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='1a1a2e', end_color='1a1a2e', fill_type='solid')
    for cell in sheet[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    
    for idx, s in enumerate(students, 1):
        p = s.profile
        sheet.append([
            idx,
            s.get_full_name(),
            s.username,
            p.student_id or '-',
            p.get_college_name(),
            p.get_department_name(),
            p.get_level_display_safe(),
            s.email or '-',
            p.phone or '-',
        ])
    
    for col in sheet.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        sheet.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="students_list.xlsx"'
    wb.save(response)
    return response


@require_role('admin')
@require_POST
def delete_account(request, pk):
    """Delete a user account."""
    user = get_object_or_404(User, pk=pk)
    if user == request.user:
        messages.error(request, 'لا يمكنك حذف حسابك الشخصي')
        return redirect('students_list')
    name = user.get_full_name() or user.username
    user.delete()
    messages.success(request, f'تم حذف حساب "{name}"')
    
    next_url = request.POST.get('next', 'students_list')
    return redirect(next_url)


@require_role('admin')
@require_POST
def reset_account_password(request, pk):
    """Reset a user's password and show the new one."""
    user = get_object_or_404(User, pk=pk)
    if user == request.user:
        messages.error(request, 'لا يمكنك تغيير كلمة مرور حسابك الشخصي من هنا')
    else:
        new_password = _generate_password()
        user.set_password(new_password)
        user.save()
        
        profile, _ = Profile.objects.get_or_create(user=user)
        profile.generated_password = new_password
        profile.save()
        
        messages.success(request, f'تم إعادة تعيين كلمة مرور الحساب "{user.get_full_name() or user.username}" بنجاح. كلمة المرور الجديدة: {new_password}')
        
    next_url = request.POST.get('next', 'account_management')
    return redirect(next_url)

def about_system(request):
    """Render the about / system info page."""
    faqs = [
        {'q': 'كيف يعمل نظام المراقبة التلقائي؟', 'a': 'يقوم النظام برصد وتخزين أي محاولات للخروج من نافذة الاختبار أو تغيير التبويبات، وسيتم إبلاغ المراقب فوراً لاتخاذ الإجراء المناسب.'},
        {'q': 'هل يمكنني مراجعة إجاباتي قبل تسليم الاختبار؟', 'a': 'نعم، في حال كان خيار المراجعة مفعلاً من قبل الأستاذ، يمكنك التنقل بين الأسئلة ومراجعتها قبل الضغط على زر "إنهاء التقديم".'},
        {'q': 'ماذا يحدث إذا انقطع الاتصال بالإنترنت أثناء الاختبار؟', 'a': 'يقوم النظام بحفظ إجاباتك تلقائياً بعد كل تفاعل بفضل ميزة الحفظ التلقائي (Auto-save)، لن تفقد إمكانية استكمال الاختبار من النقطة التي توقفت عندها.'},
    ]
    return render(request, 'exams/about_system.html', {'faqs': faqs})

@require_role(['admin', 'assistant'])
def instructors_list(request):
    from accounts.models import College, Department
    
    instructors = User.objects.filter(profile__role__in=['admin', 'assistant']).select_related(
        'profile', 'profile__college', 'profile__academic_department'
    ).order_by('last_name', 'first_name')

    college_id = request.GET.get('college', '')
    dept_id = request.GET.get('department', '')
    search_q = request.GET.get('q', '').strip()

    if college_id:
        instructors = instructors.filter(profile__college_id=college_id)
    if dept_id:
        instructors = instructors.filter(profile__academic_department_id=dept_id)
    if search_q:
        instructors = instructors.filter(
            Q(first_name__icontains=search_q) |
            Q(last_name__icontains=search_q) |
            Q(username__icontains=search_q)
        )

    colleges = College.objects.all()
    departments = Department.objects.all()
    if college_id:
        departments = departments.filter(college_id=college_id)

    total_all = User.objects.filter(profile__role__in=['admin', 'assistant']).count()

    context = {
        'instructors': instructors,
        'total': instructors.count(),
        'total_all': total_all,
        'colleges': colleges,
        'departments': departments,
        'selected_college': int(college_id) if college_id else None,
        'selected_dept': int(dept_id) if dept_id else None,
        'search_q': search_q,
    }
    return render(request, 'exams/instructors_list.html', context)

@require_role(['admin', 'assistant'])
def edit_account(request, pk):
    from accounts.models import College, Department
    user_to_edit = get_object_or_404(User, pk=pk)
    
    # Prevent assistants from editing admins (unless they are editing themselves)
    if request.user.profile.role == 'assistant' and 'admin' in user_to_edit.profile.get_allowed_roles_list() and request.user.pk != user_to_edit.pk:
        messages.error(request, 'ليس لديك صلاحية لتعديل حساب هذا المشرف.')
        return redirect('dashboard')

    if request.method == 'POST':
        # Update User
        user_to_edit.first_name = request.POST.get('first_name', user_to_edit.first_name)
        user_to_edit.last_name = request.POST.get('last_name', user_to_edit.last_name)
        user_to_edit.email = request.POST.get('email', user_to_edit.email)
        user_to_edit.save()

        # Update Profile
        profile = user_to_edit.profile
        profile.phone = request.POST.get('phone', profile.phone)
        
        college_id = request.POST.get('college')
        if college_id:
            profile.college_id = college_id
            
        dept_id = request.POST.get('department')
        if dept_id:
            profile.academic_department_id = dept_id
            
        # Update Multiple Roles
        if request.user.profile.role == 'admin': # Only admins can change allowed roles broadly
            allowed_roles = request.POST.getlist('allowed_roles')
            if allowed_roles:
                profile.allowed_roles = allowed_roles
                if profile.role not in allowed_roles:
                    profile.role = allowed_roles[0] # Fallback active role to the first allowed
            else:
                messages.warning(request, 'يجب أن يمتلك المستخدم صلاحية واحدة على الأقل. تم تجاهل تعديل الصلاحيات المفرغ.')

        if 'student' in profile.get_allowed_roles_list():
            student_id = request.POST.get('student_id')
            if student_id is not None:
                profile.student_id = student_id
            level = request.POST.get('level')
            if level:
                profile.level = int(level)
                
        profile.save()
        messages.success(request, f'تم تحديث بيانات {user_to_edit.get_full_name()} بنجاح!')
        
        next_url = request.POST.get('next', 'dashboard')
        if next_url == 'students_list':
            return redirect('students_list')
        elif next_url == 'instructors_list':
            return redirect('instructors_list')
        elif next_url == 'account_management':
            return redirect('account_management')
        return redirect('dashboard')

    context = {
        'edit_user': user_to_edit,
        'colleges': College.objects.all(),
        'departments': Department.objects.all() if not user_to_edit.profile.college else Department.objects.filter(college=user_to_edit.profile.college),
        'next_url': request.GET.get('next', 'dashboard')
    }
    return render(request, 'exams/edit_account.html', context)

@login_required
def switch_role(request):
    """Allows a user to switch their active session role if they have multiple allowed roles."""
    if request.method == 'POST':
        new_role = request.POST.get('role')
        profile = request.user.profile
        if new_role in profile.get_allowed_roles_list():
            profile.role = new_role
            profile.save()
            messages.success(request, f'تم تبديل واجهة النظام للعمل كـ: {profile.get_role_display_safe()}')
        else:
            messages.error(request, 'ليس لديك الصلاحية لاستخدام هذا الدور الدخول.')
        
    return redirect('dashboard')

@require_role(['admin', 'assistant'])
def export_exam_results_excel(request, pk):
    import openpyxl
    from django.http import HttpResponse
    
    exam = get_object_or_404(Exam, pk=pk)
    if not request.user.profile.is_admin() and exam.created_by != request.user:
        return HttpResponseForbidden()
        
    attempts = exam.attempts.filter(is_submitted=True).select_related('student')
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "نتائج الاختبار"
    
    headers = ['اسم الطالب', 'الرقم الجامعي', 'وقت البدء', 'وقت الانتهاء', 'الدرجة', 'النسبة المئوية', 'حالة النجاح']
    ws.append(headers)
    
    for attempt in attempts:
        student = attempt.student
        pct = attempt.get_percentage()
        is_passed = "ناجح" if pct >= (exam.pass_mark / exam.total_marks * 100 if exam.total_marks else 0) else "راسب"
        row = [
            student.get_full_name() or student.username,
            student.username,
            attempt.started_at.strftime('%Y-%m-%d %H:%M') if attempt.started_at else '',
            attempt.submitted_at.strftime('%Y-%m-%d %H:%M') if attempt.submitted_at else '',
            attempt.final_score,
            f"{pct}%",
            is_passed
        ]
        ws.append(row)
        
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="exam_{exam.id}_results.xlsx"'
    wb.save(response)
    return response

@require_role(['admin', 'assistant'])
def import_questions_excel(request, exam_pk):
    import openpyxl
    exam = get_object_or_404(Exam, pk=exam_pk)
    if not request.user.profile.is_admin() and exam.created_by != request.user:
        return HttpResponseForbidden()
        
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        try:
            wb = openpyxl.load_workbook(excel_file)
            ws = wb.active
            
            for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                if not row[0] or not row[1]:
                    continue
                
                qtype = str(row[0]).strip()
                qtext = str(row[1]).strip()
                marks = int(row[2]) if row[2] else 5
                
                question = Question.objects.create(
                    exam=exam,
                    question_type=qtype,
                    text=qtext,
                    marks=marks,
                    order=exam.questions.count() + 1
                )
                
                if qtype == 'mcq_single':
                    options = [row[3], row[4], row[5], row[6]]
                    correct_idx = int(row[7]) if row[7] else 1
                    for i, opt in enumerate(options):
                        if opt:
                            QuestionOption.objects.create(
                                question=question,
                                text=str(opt).strip(),
                                is_correct=(i + 1 == correct_idx),
                                order=i
                            )
                elif qtype == 'true_false':
                    correct_ans = str(row[7]).strip().lower() in ['true', '1', 'صح', 'صحيح']
                    question.tf_answer = correct_ans
                    question.save()
                    
            messages.success(request, 'تم استيراد الأسئلة بنجاح')
        except Exception as e:
            messages.error(request, f'حدث خطأ أثناء الاستيراد: {str(e)}')
            
        return redirect('exam_edit', pk=exam.pk)
    
    return render(request, 'exams/import_questions.html', {'exam': exam})
